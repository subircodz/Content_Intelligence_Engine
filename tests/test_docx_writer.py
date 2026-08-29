import re
from pathlib import Path

from docx import Document

from intelligence_content_engine.output.docx_writer import (
    markdown_to_docx,
    safe_filename,
    save_article_docx,
)


def test_safe_filename_generation() -> None:
    assert (
        safe_filename("How We Evaluate Online Casinos: Power.win Editorial & Review Methodology")
        == "how-we-evaluate-online-casinos-power-win-editorial-review-methodology"
    )


def test_safe_filename_handles_special_chars() -> None:
    assert safe_filename("Hello!!! World???") == "hello-world"
    assert safe_filename("   ") == "untitled-article"
    assert safe_filename("A/B Test") == "a-b-test"


def test_docx_file_created(tmp_path: Path) -> None:
    article = "# Heading\n\nSome paragraph text."
    path = save_article_docx(article, "Test Title", output_dir=str(tmp_path))

    assert path is not None
    assert path.exists()
    assert path.suffix == ".docx"


def test_empty_article_does_not_create_docx(tmp_path: Path) -> None:
    assert save_article_docx("", "Test Title", output_dir=str(tmp_path)) is None
    assert save_article_docx("   ", "Test Title", output_dir=str(tmp_path)) is None
    assert list(tmp_path.iterdir()) == []


def test_user_title_preserved_as_document_title(tmp_path: Path) -> None:
    user_title = "How We Evaluate Online Casinos"
    article = "Body text."
    doc = markdown_to_docx(article, user_title)

    first_paragraph = doc.paragraphs[0]
    assert first_paragraph.text == user_title
    assert "Title" in first_paragraph.style.name


def test_markdown_headings_converted(tmp_path: Path) -> None:
    article = "# Main Heading\n\nIntro.\n\n## Sub Heading\n\nMore text."
    doc = markdown_to_docx(article, "Doc Title")

    heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "Main Heading" in heading_texts
    assert "Sub Heading" in heading_texts


def test_bullet_and_numbered_lists_converted() -> None:
    article = "- First item\n- Second item\n\n1. One\n2. Two"
    doc = markdown_to_docx(article, "List Doc")

    styles = [p.style.name for p in doc.paragraphs]
    assert styles.count("List Bullet") == 2
    assert styles.count("List Number") == 2


def test_bold_italic_preserved_no_raw_markdown() -> None:
    article = "This has **bold** and *italic* text."
    doc = markdown_to_docx(article, "Format Doc")

    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "**" not in body_text
    # Find the paragraph with formatting
    fmt_para = next(p for p in doc.paragraphs if "bold" in p.text)
    bold_runs = [r for r in fmt_para.runs if r.bold]
    italic_runs = [r for r in fmt_para.runs if r.italic]
    assert len(bold_runs) > 0
    assert len(italic_runs) > 0


def test_docx_generation_error_returns_none(tmp_path: Path) -> None:
    # Force an error by passing an invalid title type scenario via monkeypatching Document.save
    from unittest.mock import patch

    with patch("intelligence_content_engine.output.docx_writer.Document") as mock_doc_cls:
        mock_doc_cls.return_value.save.side_effect = OSError("disk full")
        result = save_article_docx("Article", "Title", output_dir=str(tmp_path))
        assert result is None
